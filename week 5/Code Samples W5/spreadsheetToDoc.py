from openpyxl import load_workbook
from docx import Document

# Open an XSLX file (remember these can contain multiple spreadsheets!)
workbook = load_workbook(filename="spread2doc.xlsx")

# List of sheet names in the workbook
print(workbook.sheetnames)

# Select the active spreadsheet (the one that is displayed when the file opens)
sheet = workbook.active

# Get a sheet's title
print(sheet.title)

# Access cell data in a sheet
print(sheet['A1'].value)

records = []

# Convert sheet data into records
for i in range(1,6) :
    record = [sheet[f'A{i}'].value,\
        sheet[f'B{i}'].value,\
        sheet[f'C{i}'].value]
    records.append(record)
print(records)

# Create a word document
doc = Document()

# Add the sheet name as a heading
doc.add_heading(sheet.title, 0)

# Create/populate the table
tbl = doc.add_table(rows=0, cols=3)
for record in records:
    new_row = tbl.add_row().cells
    new_row[0].text = record[0]
    new_row[1].text = str(record[1])
    new_row[2].text = record[2]

# Make header cells bold
for cell in tbl.rows[0].cells:
    for p in cell.paragraphs:
        for run in p.runs:
            run.font.bold = True

doc.save("spread2doc.docx")

# Example from https://openpyxl.readthedocs.io/en/stable/
from openpyxl import Workbook

wb = Workbook()

st = wb.active

# Edit a cell
st['A1'] = "Test"

# We can also use built-in Python types and they will be converted
import datetime
st['A2'] = datetime.datetime.now()

# Append a row
st.append(["ColA","ColB","ColC"])

# Save the sheet
wb.save("newsheet.xlsx")