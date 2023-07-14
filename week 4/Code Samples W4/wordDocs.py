from docx import Document
from docx.shared import Inches

"""
Creating/Saving a word document
"""

# Create a new document
document = Document()

# Add a header
document.add_heading('Document Title', 0)

# Create a paragraph with some text initially
p = document.add_paragraph('A plain paragraph having some ')

# Add text to the paragraph with different properties
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

# Add more styled headings and paragraphs
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

# Add bullet points
document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)

# Add a picture
document.add_picture('royaltyFree0.gif', width=Inches(1.25))

# Create table data
records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

# Create a table and fill it with data
table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

# Add a page break
document.add_page_break()

# Save the created document to a file named demo.docx
document.save('demo.docx')

"""
Opening/Editing an existing document
"""

# Open existing document
existingDoc = Document("demo.docx")

# View document XML (Helps with understanding navigation)
print(existingDoc._element.xml)

# Access and edit paragraphs
for i in range(len(existingDoc.paragraphs)):
    print(f"P{i}: {existingDoc.paragraphs[i].text}")

existingDoc.paragraphs[0].text = "A New Title"
existingDoc.paragraphs[1].text = "Some new body text"

# Print edited document
for i in range(len(existingDoc.paragraphs)):
    print(f"P{i}: {existingDoc.paragraphs[i].text}")

"""
Simple search and replace function
"""

def replaceAll(old, new, doc):
    """ Replaces all instances of a string in a word document """
    for p in doc.paragraphs:
        p.text = str.replace(p.text, old, new)

replaceAll("Some", "Lots of", existingDoc)

# Print edited document
for i in range(len(existingDoc.paragraphs)):
    print(f"P{i}: {existingDoc.paragraphs[i].text}")

# Save the document
existingDoc.save("demo.docx")