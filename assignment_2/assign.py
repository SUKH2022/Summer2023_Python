import os
from docx import Document
from openpyxl import Workbook

def processInvoices():
    # Initialize
    workbook = Workbook()
    sheet = workbook.active

    # Add column headers in the word doc
    headers = ['Invoice ID', 'Total Quantity', 'Subtotal', 'Tax', 'Total']
    sheet.append(headers)

    # Get the list of doc files
    files = [file for file in os.listdir('.') if file.endswith('.docx')]

    # Process each doc file
    for file in files:
        doc = Document(file)

        # Extract invoice details from the documents
        invoice_id = doc.paragraphs[0].text[1:]
        products = doc.paragraphs[2:-3]
        total_quantity = sum([int(product.text.split(':')[1]) for product in products])
        subtotal = float(doc.paragraphs[-3].text.split(':')[1])
        tax = float(doc.paragraphs[-2].text.split(':')[1])
        total = float(doc.paragraphs[-1].text.split(':')[1])

        # Add invoice details to the word doc
        row = [invoice_id, total_quantity, subtotal, tax, total]
        sheet.append(row)

    # Save the the word doc
    workbook.save('InvoiceSummary.xlsx')

# Run the function to process the invoices
processInvoices()